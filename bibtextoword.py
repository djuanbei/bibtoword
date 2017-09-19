#!/usr/bin/python
'''
Created on Sep. 19, 2017

@author: dailiyun
'''



import os
import sys
from docx import Document
from docx.shared import Inches
import bibtexparser

group_author=['Joon'] 
must_element=['author','year','title','pages' ]

class bibbase():
    def __init__(self,filename):
        self.perYearPublish={}
        with open('bibtex.bib') as bibtex_file:
            self.bibtex_str = bibtex_file.read()
        bib_database = bibtexparser.loads(self.bibtex_str)
        listterms=bib_database.get_entry_list()
        self.num=0
        for e in listterms:
            haveTotalMustElem=True
            for ele in must_element:
                if ele not in e:
                    haveTotalMustElem=False
            if haveTotalMustElem:
                self.num+=1
                yy=e['year']
                if yy in self.perYearPublish:
                    self.perYearPublish[yy].append(e)
                else:
                    self.perYearPublish[yy]=[e]

    def addAuthour(self, para, paper):
        authors=paper['author']
        authorList=authors.split(',')
        first=True
        for authors  in authorList:
            authorsecondLish=authors.split("and")
            for author in authorsecondLish:
                author=author.strip()
                if not first:
                    para.add_run(", ")
                first=False
                
                if author in group_author:
                    para.add_run(author).bold = True
                else:
                    para.add_run(author)
                    
                    
        para.add_run(': ')
        
    def addOneYear(self, document,y):
        document.add_heading(str(y), level=1)

        papers=self.perYearPublish[y]
        for paper in papers:

            para = document.add_paragraph('['+str(self.num)+']. ')
            self.addAuthour(para,paper)
            
            para.add_run(paper['title']+'. ').bold = True
            if  'journal' in paper :
                para.add_run(paper['journal']+", ").italic = True
            elif 'booktitle' in paper:
                para.add_run(paper['booktitle']+", ").italic = True

            if 'volume' in paper :
                para.add_run(paper['volume'])
                if 'issn' in paper:
                    para.add_run("("+paper['issn']+")")

            para.add_run(": "+paper['pages']+" ("+str(y)+")")
            self.num-=1
            
    def ouputdoc(self, outname):
        document = Document()
        document.add_heading('RISE Publications', 0)
        years=self.perYearPublish.keys()
        years.sort()
        years.reverse()
        for y in years:
            self.addOneYear(document,y)

        document.save(outname)
        
    
if __name__ == "__main__":
    bib=bibbase(sys.argv[1])
    bib.ouputdoc(sys.argv[2])
    
